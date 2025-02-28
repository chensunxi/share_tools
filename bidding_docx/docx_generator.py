import shutil
import sys
import os

import fitz  # PyMuPDF
from typing import Dict, List, Tuple
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Mm, Cm, Inches
from PIL import Image
# from utils.ocr_extractor_utils import OcrExtractorUtils
from utils.logger_utils import LoggerUtils
from utils.pdf_process_utils import PdfProcessUtils
from utils.common_utils import extract_chinese


class EmployeeDocumentGenerator:
    def __init__(self, template_path: str, out_logger):
        """
        初始化文档生成器
        :param template_path: Word模板文件路径
        """
        self.logger = LoggerUtils.get_logger(None, out_logger)
        # self.extractor = OcrExtractorUtils()
        self.real_name = None

        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")

        try:
            self.doc = DocxTemplate(template_path)
            self.logger.info(f"Template loaded successfully: {template_path}")
        except Exception as e:
            self.logger.error(f"Error loading template: {str(e)}")
            raise

    def _get_image_orientation(self, image_path: str) -> str:
        """
        获取图片方向
        :param image_path: 图片路径
        :return: 'landscape' 或 'portrait'
        """
        try:
            with Image.open(image_path) as img:
                width, height = img.size
                img.verify()  # 验证文件是否损坏
                return 'landscape' if width > height else 'portrait'
        except Exception as e:
            self.logger.error(f"Error getting image orientation for {image_path}: {str(e)}")
            return 'unknown'

    def _get_image_size(self, image_path: str) -> Tuple[int, int]:
        """
        获取图片尺寸
        :param image_path: 图片路径
        :return: (宽度, 高度)
        """
        try:
            with Image.open(image_path) as img:
                width, height = img.size
                return width, height
        except Exception as e:
            self.logger.error(f"Error getting image size for {image_path}: {str(e)}")
            return (0, 0)

    def _calculate_image_dimensions(self, width: int, height: int, target_width: int = 160) -> Tuple[int, int]:
        """
        计算保持纵横比的图片尺寸
        :param width: 原始宽度
        :param height: 原始高度
        :param target_width: 目标宽度(mm)
        :return: (新宽度, 新高度) in mm
        """
        if width == 0 or height == 0:
            return target_width, target_width  # 默认正方形
        aspect_ratio = width / height
        new_width = target_width
        new_height = int(target_width / aspect_ratio)
        return new_width, new_height

    def _prepare_image(self, image_path: str, image_type: str, output_dir: str) -> Tuple[str, str]:
        """
        处理图片文件
        :return: (处理后的图片路径, 图片方向)
        """
        try:
            if not os.path.exists(image_path):
                raise FileNotFoundError(f"Image file not found: {image_path}")

            _, ext = os.path.splitext(image_path)
            output_path = os.path.join(output_dir, f"{image_type}{ext}")

            shutil.copy2(image_path, output_path)
            orientation = self._get_image_orientation(output_path)

            self.logger.info(f"Image processed: {output_path}, orientation: {orientation}")
            return output_path, orientation

        except Exception as e:
            self.logger.error(f"Error processing image {image_path}: {str(e)}")
            raise

    def _convert_pdf_to_images(self, pdf_path: str, file_type: str, output_dir: str) -> List[str]:
        """
        将PDF转换为图片，截取部分页面
        :param pdf_path: PDF文件路径
        :param output_dir: 输出目录
        :return: 生成的图片路径列表
        """
        # 设置DPI来提高图像分辨率
        dpi = 300  # 可根据需求调整
        try:
            if not os.path.exists(pdf_path):
                raise FileNotFoundError(f"PDF file not found: {pdf_path}")

            keywords = ["身份证号码", "甲方盖章："]
            pdf_document = fitz.open(pdf_path)
            total_pages = len(pdf_document)
            self.logger.info(f"Converting PDF {pdf_path}: total pages = {total_pages}")

            image_paths = []

            # 确定要处理的页面
            pages_to_convert = []
            if file_type == 'labor_contract':
                # 如果是劳动合同，根据关键字查找3个页面（封面、甲乙双方信息页、劳动合同签字盖章页)
                _, found_pages = PdfProcessUtils.search_keywords_in_pdf(pdf_document, keywords)
                pages_to_convert = [0] + found_pages  # 合并第一页和找到关键字的页面
            else:
                # 其他PDF文件，处理所有页面
                pages_to_convert = range(total_pages)


            # 转换选定的页面
            for page_num in pages_to_convert:
                page = pdf_document[page_num]
                if file_type == 'social_security_proof':
                    PdfProcessUtils.draw_rect_name(page, self.real_name)
                    PdfProcessUtils.draw_rect_unit(page)
                pix = page.get_pixmap(dpi=dpi)

                output_image_path = os.path.join(output_dir, f"{file_type}_page_{page_num}.png")

                pix.save(output_image_path)
                image_paths.append(output_image_path)
                self.logger.info(f"Converted page {page_num} to {output_image_path}")

            pdf_document.close()
            return image_paths

        except Exception as e:
            self.logger.error(f"Error converting PDF {pdf_path}: {str(e)}")
            raise

    def optimize_cert_order(self, cert_items):
        """优化证书顺序，使横向证书尽量配对"""
        # 按方向分组
        horizontal_certs = []
        vertical_certs = []

        for cert in cert_items:
            if self._get_image_orientation(cert['path']) == 'landscape':
                horizontal_certs.append(cert)
            else:
                vertical_certs.append(cert)

        # 重新排序
        optimized_certs = []

        # 先处理横向证书（两两配对）
        for i in range(0, len(horizontal_certs), 2):
            optimized_certs.append(horizontal_certs[i])
            if i + 1 < len(horizontal_certs):
                optimized_certs.append(horizontal_certs[i + 1])

        # 添加纵向证书
        optimized_certs.extend(vertical_certs)

        # 输出优化顺序
        self.logger.info("证书优化顺序：")
        for i, cert in enumerate(optimized_certs, 1):
            direction = "横向" if self._get_image_orientation(cert['path']) == 'landscape' else "纵向"
            self.logger.info(f"{i}. {cert['name']} ({direction}): {cert['path']}")

        return optimized_certs

    def prepare_employee_data(self, emp_data: Dict, temp_dir: str, emp_idx: int) -> Dict:
        """
        准备单个员工的数据
        """
        """
            准备单个员工的数据，增加对可能为 None 的字段的兼容处理
            """
        try:
            # 创建临时目录
            employee_temp_dir = os.path.join(temp_dir, f'employee_{emp_idx}')
            os.makedirs(employee_temp_dir, exist_ok=True)

            # 获取图片尺寸
            def get_image_size_and_display(image_path: str, orientation: str, landscape_width, portrait_width) -> Tuple[int, int, int, int]:
                """ 获取图片尺寸和显示尺寸 """
                if not image_path:
                    return 0, 0, 0, 0
                try:
                    width, height = self._get_image_size(image_path)
                    target_width = landscape_width if orientation == 'landscape' else portrait_width
                    display_width, display_height = self._calculate_image_dimensions(width, height, target_width)
                    return width, height, display_width, display_height

                except Exception as e:
                    self.logger.error(f"Error get_image_size_and_display for {image_path}: {str(e)}")
                return 0, 0, 0, 0

            # 处理各类图片，增加兼容性检查
            def safe_prepare_image_field(field_name: str, image_type: str) -> Tuple[str, str, InlineImage]:
                """ 处理图片字段，返回图片路径和方向信息，若无数据则返回 None """
                image_path = emp_data.get(field_name)
                if image_path:
                    try:
                        image_path, orientation = self._prepare_image(image_path, image_type, employee_temp_dir)
                        if orientation == 'unknown':
                            return None, None, None
                        if image_type == 'id_card':
                            return image_path, orientation, InlineImage(self.doc, image_path, width=Mm(77.9))
                        elif image_type in ['graduation_cert', 'degree_cert']:
                            width, height, display_width, display_height = get_image_size_and_display(image_path, orientation, 128, 143)
                            return image_path, orientation, InlineImage(self.doc, image_path, width=Mm(display_width), height=Mm(display_height))
                        else:
                            width, height, display_width, display_height = get_image_size_and_display(image_path, orientation, 160, 171)
                            return image_path, orientation, InlineImage(self.doc, image_path, width=Mm(display_width), height=Mm(display_height))
                    except Exception as e:
                        self.logger.error(f"Error safe_prepare_image_field for {image_path}: {str(e)}")
                        return None, None, None

                return None, None, None

            # 获取员工姓名
            self.real_name = extract_chinese(emp_data.get('name', ''))

            # 获取学信网图片路径
            edu_screenshot_path = emp_data.get('china_edu_screenshot', None)
            # 提取教育背景信息
            school_name, major_name = None, None
            # if edu_screenshot_path:
            #     result = self.extractor.extract_info(edu_screenshot_path)
            #     if result['success']:
            #         extracted_info = result['data']
            #         school_name = extracted_info.get('学校名称')
            #         major_name = extracted_info.get('专业')

            # 准备基本数据，避免 None 值的字段会抛出错误
            prepared_data = {
                'index': emp_idx,
                'name': self.real_name,
                'id_no': emp_data.get('id_no', ''),
                'school': school_name,
                'major': major_name,
                'cert_orientations': {
                    'graduation': None,
                    'degree': None
                }
            }

            # 处理学信网截图
            if edu_screenshot_path:
                if edu_screenshot_path.lower().endswith('.pdf'):
                    edu_screenshot_page = self._convert_pdf_to_images(edu_screenshot_path, 'china_edu_screenshot',
                                                                       employee_temp_dir)
                    prepared_data['china_edu_screenshot'] = InlineImage(self.doc, edu_screenshot_page[0], width=Mm(160))
                else:
                    _, _, edu_screenshot_img = safe_prepare_image_field('china_edu_screenshot', 'china_edu')
                    prepared_data['china_edu_screenshot'] = edu_screenshot_img
            else:
                prepared_data['china_edu_screenshot'] = None

            # 处理学历证书
            graduation_cert_path = emp_data.get('graduation_cert', None)
            if graduation_cert_path:
                if graduation_cert_path.lower().endswith('.pdf'):
                    graduation_cert_page = self._convert_pdf_to_images(graduation_cert_path, 'graduation_cert', employee_temp_dir)
                    prepared_data['graduation_cert'] = InlineImage(self.doc, graduation_cert_page[0], width=Mm(128))
                    grad_orientation = self._get_image_orientation(graduation_cert_page[0])
                else:
                    grad_cert_path, grad_orientation, grad_cert_img = safe_prepare_image_field('graduation_cert', 'graduation_cert')
                    prepared_data['graduation_cert'] = grad_cert_img
                prepared_data['cert_orientations']['graduation'] = grad_orientation
            else:
                prepared_data['graduation_cert'] = None

            # 处理学位证书
            degree_cert_path = emp_data.get('degree_cert', None)
            if degree_cert_path:
                if degree_cert_path.lower().endswith('.pdf'):
                    degree_cert_page = self._convert_pdf_to_images(degree_cert_path, 'degree_cert',
                                                                       employee_temp_dir)
                    prepared_data['degree_cert'] = InlineImage(self.doc, degree_cert_page[0], width=Mm(128))
                    degree_orientation = self._get_image_orientation(degree_cert_page[0])
                else:
                    degree_cert_path, degree_orientation, degree_cert_img = safe_prepare_image_field('degree_cert', 'degree_cert')
                    prepared_data['degree_cert'] = degree_cert_img
                prepared_data['cert_orientations']['degree'] = degree_orientation
            else:
                prepared_data['degree_cert'] = None

            # 处理身份证
            id_card_front = emp_data.get('id_card_front', None)
            if id_card_front:
                if id_card_front.lower().endswith('.pdf'):
                    id_card_front_page = self._convert_pdf_to_images(id_card_front, 'id_card', employee_temp_dir)
                    # id_card_back_img = self._convert_pdf_to_images(id_card_front, 'id_card', employee_temp_dir)
                    prepared_data['id_card_front'] = InlineImage(self.doc, id_card_front_page[0], width=Mm(77.9))
                    prepared_data['id_card_back'] = InlineImage(self.doc, id_card_front_page[0], width=Mm(77.9))
                else:
                    _, _, id_card_front_img = safe_prepare_image_field('id_card_front','id_card')
                    # _, _, id_card_back_img = safe_prepare_image_field('id_card_back', 'id_card')
                    prepared_data['id_card_front'] = id_card_front_img
                    prepared_data['id_card_back'] = id_card_front_img
            else:
                prepared_data['id_card_front'] = None
                prepared_data['id_card_back'] = None

            # 处理劳动合同
            labor_contract = emp_data.get('labor_contract', None)
            if labor_contract:
                labor_contract_images = self._convert_pdf_to_images(labor_contract, 'labor_contract',
                                                                employee_temp_dir)
                prepared_data['labor_contract'] = [
                    InlineImage(self.doc, img_path, width=Mm(151.5))
                    for img_path in labor_contract_images if img_path
                ]
            else:
                prepared_data['labor_contract'] = []

            # 处理社保证明
            social_security_proof = emp_data.get('social_security_proof', None)
            if social_security_proof:
                social_security_images = self._convert_pdf_to_images(
                    social_security_proof, 'social_security_proof', employee_temp_dir
                )
                prepared_data['social_security_proof'] = [
                    InlineImage(self.doc, img_path, width=Mm(151))
                    for img_path in social_security_images
                ]
            else:
                prepared_data['social_security_proof'] = []

            # 处理资格证书
            qualification_certs = emp_data.get('qualification_certs', None)
            prepared_data['qualification_certs'] = []

            if qualification_certs:
                # 准备证书数据
                cert_items = [{'name': name, 'path': path}
                              for name, path in qualification_certs.items()]

                # 优化证书顺序
                optimized_certs = self.optimize_cert_order(cert_items)

                for cert in optimized_certs:
                    cert_image_path = cert.get('path', None)
                    if cert_image_path.lower().endswith('.pdf'):
                        cert_image_page = self._convert_pdf_to_images(cert_image_path, 'qual_cert', employee_temp_dir)
                        cert_targe_path = cert_image_page[0]
                        cert_orientation = self._get_image_orientation(cert_targe_path)
                    else:
                        cert_targe_path, cert_orientation = self._prepare_image(cert['path'], f'qual_cert_{cert["name"]}', employee_temp_dir)

                    cert_width, cert_height, cert_display_width, cert_display_height = get_image_size_and_display(
                        cert_targe_path, cert_orientation,128, 143)
                    cert_image = InlineImage(self.doc, cert_targe_path, width=Mm(cert_display_width),
                                           height=Mm(cert_display_height))

                    prepared_data['qualification_certs'].append({
                        'name': cert['name'],
                        'img': cert_image,
                        'orientation': cert_orientation
                    })

            return prepared_data

        except Exception as e:
            self.logger.error(f"Error preparing employee data: {str(e)}")
            raise

    def generate_document(self, employees_data: List[Dict], output_path: str):
        """
        生成文档
        """
        try:
            self.logger.info("Starting document generation...")

            output_dir = os.path.dirname(output_path)
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)

            temp_dir = os.path.join(output_dir, 'temp_images')
            os.makedirs(temp_dir, exist_ok=True)

            prepared_employees = []
            for idx, emp_data in enumerate(employees_data, 1):
                self.logger.info(f"Processing employee {idx}: {emp_data['name']}")
                prepared_emp_data = self.prepare_employee_data(emp_data, temp_dir, idx)
                prepared_employees.append(prepared_emp_data)
                self.logger.info(f"Employee {emp_data['name']} data prepared successfully")

            context = {'employees': prepared_employees}
            self.logger.info("Rendering template...")
            print(f"渲染上下文:{context}")
            # 生成临时文档
            temp_output = os.path.join(output_dir, 'temp_output.docx')
            try:
                self.doc.render(context)
            except Exception as e:
                self.logger.error(f"Error rendering template: {e}")
                self.logger.error(traceback.format_exc())
            self.doc.save(temp_output)
            # 处理分页
            doc = Document(temp_output)
            previous_para_text = ""

            # 从段落文本中提取员工索引
            emp_idx = None
            for para in doc.paragraphs:
                has_image = False
                for run in para.runs:
                    if len(run._element.xpath('.//w:drawing')) > 0:
                        has_image = True
                        break
                    else:
                        for emp in prepared_employees:
                            if f"{emp['name']}" in para.text:
                                emp_idx = emp['index']
                                # 跳过第一个员工
                                if emp['index'] == 1:
                                    continue
                                # 保存原始文本
                                original_text = para.text
                                # 清空原段落
                                para.clear()
                                # 先添加分页符
                                run = para.add_run()
                                run.add_break(WD_BREAK.PAGE)
                                # 再添加原始文本
                                para.add_run(original_text)
                                self.logger.info(f"在员工 {emp_data['name']} 章节前添加分页符")

                if has_image:
                    cert_type = None

                    if '毕业证' in previous_para_text:
                        cert_type = 'graduation'
                    elif '学位证' in previous_para_text:
                        cert_type = 'degree'

                    # 处理毕业证和学位证的分页逻辑
                    if cert_type and emp_idx is not None:
                        emp_data = prepared_employees[emp_idx - 1]
                        orientations = emp_data['cert_orientations']

                        # 只有当证书不是都是横向时才添加分页符
                        if cert_type == 'graduation' and (
                                orientations['graduation'] != 'landscape' or
                                orientations['degree'] != 'landscape'
                        ):
                            run = para.add_run()
                            run.add_break(WD_BREAK.PAGE)
                            self.logger.info(f"在员工{emp_idx}: {emp_data['name']}的毕业证书后添加分页符")

                    # 处理资质证书的分页
                    emp_data = prepared_employees[emp_idx - 1]
                    for cert_idx, cert in enumerate(emp_data['qualification_certs']):
                        # 直接查找证书名称
                        if cert['name'] in previous_para_text:
                            if cert_idx < len(emp_data['qualification_certs']) - 1:
                                next_cert = emp_data['qualification_certs'][cert_idx + 1]
                                if not (cert['orientation'] == 'landscape' and next_cert['orientation'] == 'landscape' and cert_idx % 2 == 0):
                                    run.add_break(WD_BREAK.PAGE)
                                    self.logger.info(f"在证书 {cert['name']} 后添加分页符")
                    # 其他文档的分页处理
                    # elif '身份证' in previous_para_text or '社保' in previous_para_text:
                    #     run = para.add_run()
                    #     run.add_break(WD_BREAK.PAGE)

                previous_para_text = para.text

            # 保存最终文档
            self.logger.info(f"Saving document to: {output_path}")
            doc.save(output_path)

            # 清理临时文件
            os.remove(temp_output)
            shutil.rmtree(temp_dir, ignore_errors=True)
            self.logger.info("Temporary files cleaned up")
            self.logger.info("Document generated successfully")

        except Exception as e:
            self.logger.error(f"Error generating document: {str(e)}")
            self.logger.error("Stack trace:", exc_info=True)
            raise


if __name__ == '__main__':
    # 示例数据
    # employees_data = [
    #     {
    #         'name': '张三',
    #         'id_card_front': r'D:\project\share_tools\bidding_docx\files\39887\id_card.jpg',
    #         'id_card_back': r'D:\project\share_tools\bidding_docx\files\39887\id_card.jpg',
    #         'graduation_cert': r'D:\project\share_tools\bidding_docx\files\39887\graduation_cert.png',
    #         'degree_cert': r'D:\project\share_tools\bidding_docx\files\39887\degree_cert.png',
    #         'china_edu_screenshot': r'D:\project\share_tools\bidding_docx\files\39887\china_edu.png',
    #         'labor_contract': r'D:\project\share_tools\bidding_docx\files\39887\labor_contract.pdf',
    #         'qualification_certs': {
    #             '高级工程师': r'D:\project\share_tools\bidding_docx\files\39887\qual_cert1.png',
    #             'PMP证书': r'D:\project\share_tools\bidding_docx\files\39887\qual_cert3.png'
    #         },
    #         'social_security_proof': r'D:\project\share_tools\bidding_docx\files\39887\social_security.pdf'
    #     },
    #     {
    #         'name': '左泉12',
    #         'id_card_front': r'D:\project\share_tools\bidding_docx\files\39888\id_card.jpg',
    #         'id_card_back': r'D:\project\share_tools\bidding_docx\files\39888\id_card.jpg',
    #         'graduation_cert': r'D:\project\share_tools\bidding_docx\files\39888\graduation_cert.png',
    #         'degree_cert': r'D:\project\share_tools\bidding_docx\files\39888\degree_cert.png',
    #         'china_edu_screenshot': r'D:\project\share_tools\bidding_docx\files\39888\china_edu.png',
    #         'labor_contract': r'D:\project\share_tools\bidding_docx\files\39888\labor_contract.pdf',
    #         'qualification_certs': {
    #             '系统分析师': r'D:\project\share_tools\bidding_docx\files\39888\qual_cert1.png',
    #             'CISSP证书': r'D:\project\share_tools\bidding_docx\files\39888\qual_cert2.png'
    #         },
    #         'social_security_proof': r'D:\project\share_tools\bidding_docx\files\39888\social_security.pdf'
    #     }
    # ]

    employees_data = [{
        'name': '曹立丹',
        'id_no': '3502199012012321',
        'id_card_front': 'D:\\project\\share_tools\\bidding_docx\\test\\曹立丹_社保证明.pdf',
        'id_card_back': None,
        'graduation_cert':
        'D:\\project\\share_tools\\download\\20241231155940\\39887_曹立丹_学历证书.png',
        'degree_cert': None,
        'china_edu_screenshot': 'D:\\project\\share_tools\\download\\20241231155940\\39887_曹立丹_学历验证报告.png',
        'labor_contract': 'D:\\project\\share_tools\\download\\20241231155940\\39887_曹立丹_劳动合同.pdf',
        'qualification_certs': {'PMP': 'D:\\project\\share_tools\\download\\20241231155940\\39887_曹立丹_PMP.png',
        'PMI-ACP': 'D:\\project\\share_tools\\download\\20241231155940\\39887_曹立丹_PMI-ACP.png'},
        'social_security_proof': 'D:\\project\\share_tools\\bidding_docx\\test\\曹立丹_社保证明.pdf'
    },
    {
        'name': '曹俊杰',
        'id_no': '3502199503063352',
        'id_card_front': 'D:\\project\\share_tools\\download\\20241231155940\\39888_曹俊杰_身份证.jpg',
        'id_card_back': None,
        'graduation_cert': None,
        'degree_cert': 'D:\\project\\share_tools\\download\\20241231155940\\39888_曹俊杰_学位证书.jpeg',
        'china_edu_screenshot': 'D:\\project\\share_tools\\download\\20241231155940\\39888_曹俊杰_学历验证报告.jpeg',
        'labor_contract': None,
        'qualification_certs': None,
        'social_security_proof': None
    }]

    # employees_data = [{
    #     'name': '李巧玲',
    #     'id_no': '450321198511106041',
    #     'id_card_front': 'D:\\project\\share_tools\\download\\20241231184411\\01582_李巧玲_身份证.pdf',
    #     'id_card_back': None,
    #     'graduation_cert': 'D:\\project\\share_tools\\download\\20241231184411\\01582_李巧玲_学历证书.jpeg',
    #     'degree_cert': 'D:\\project\\share_tools\\download\\20241231184411\\01582_李巧玲_学位证书.jpeg',
    #     'china_edu_screenshot': None,
    #     'labor_contract': 'D:\\project\\share_tools\\download\\20241231184411\\01582_李巧玲_劳动合同.pdf',
    #     'qualification_certs': {
    #         '计算机等级考试证书': 'D:\\project\\share_tools\\download\\20241231184411\\01582_李巧玲_计算机等级考试证书.png'},
    #     'social_security_proof': None
    # },
    #     {
    #         'name': '王秀丽',
    #         'id_no': '341222199002245984',
    #         'id_card_front': 'D:\\project\\share_tools\\download\\20241231184411\\04797_王秀丽_身份证.pdf',
    #         'id_card_back': None,
    #         'graduation_cert': 'D:\\project\\share_tools\\download\\20241231184411\\04797_王秀丽_学历证书.pdf',
    #         'degree_cert': 'D:\\project\\share_tools\\download\\20241231184411\\04797_王秀丽_学位证书.pdf',
    #         'china_edu_screenshot': None,
    #         'labor_contract': None,
    #         'qualification_certs': {},
    #         'social_security_proof': None
    #     }]
    # employees_data = [{
    #     'name': '李巧玲',
    #     'id_no': '450321198511106041',
    #     'id_card_front': 'D:\\project\\share_tools\\download\\20241231184411\\01582_李巧玲_身份证.pdf',
    #     'id_card_back': None,
    #     'graduation_cert': 'D:\\project\\share_tools\\download\\20241231184411\\01582_李巧玲_学历证书.jpeg',
    #     'degree_cert': 'D:\\project\\share_tools\\download\\20241231184411\\01582_李巧玲_学位证书.jpeg',
    #     'china_edu_screenshot': None,
    #     'labor_contract': 'D:\\project\\share_tools\\download\\20241231184411\\01582_李巧玲_劳动合同.pdf',
    #     'qualification_certs': {
    #         '计算机等级考试证书': 'D:\\project\\share_tools\\download\\20241231184411\\04797_王秀丽_学历证书.pdf',
    #         'PMP证书': 'D:\\project\\share_tools\\download\\20241231184411\\01582_李巧玲_学位证书.jpeg'},
    #     'social_security_proof': None
    # },
    #     {
    #         'name': '王秀丽',
    #         'id_no': '341222199002245984',
    #         'id_card_front': 'D:\\project\\share_tools\\download\\20241231184411\\04797_王秀丽_身份证.pdf',
    #         'id_card_back': None,
    #         'graduation_cert': 'D:\\project\\share_tools\\download\\20241231184411\\04797_王秀丽_学历证书.pdf',
    #         'degree_cert': 'D:\\project\\share_tools\\download\\20241231184411\\04797_王秀丽_学位证书.pdf',
    #         'china_edu_screenshot': None,
    #         'labor_contract': None,
    #         'qualification_certs': {},
    #         'social_security_proof': None
    #     }]


    try:
        template_path = r'D:\project\share_tools\bidding_docx\test\bidding_template1.docx'
        output_path = r'D:\project\share_tools\bidding_docx\test\employee_documents.docx'

        generator = EmployeeDocumentGenerator(template_path, None)
        generator.generate_document(employees_data, output_path)

    except Exception as e:
        print(f"Error: {str(e)}")
        sys.exit(1)
