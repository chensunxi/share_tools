from datetime import datetime
import gc
import os
from io import BytesIO
import re
import pandas as pd
from docx.oxml.ns import qn

import fitz
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from PIL import Image
from utils.logger_utils import LoggerUtils
from utils.common_utils import find_pdf_files,extract_chinese
from utils.pdf_process_utils import PdfProcessUtils

class SocialDocument():
    def __init__(self, pdf_dir, excel_file, out_logger):
        self.pdf_dir = pdf_dir
        self.excel_file = excel_file
        self.out_logger = out_logger
        self.logger = LoggerUtils.get_logger(None, out_logger)


    # 将社保插入到Word 的转换
    def insert_socialdocx_to_word(self):
        # 创建一个新的 Word 文档对象
        doc = Document()

        # 获取要处理的所有pdf文件
        match_files = []
        name_pdf_mapping = {}
        employee_names = pd.read_excel(self.excel_file, skiprows=1)['工作名'].tolist()
        # 2. 按姓名查找相关PDF文件
        for employee_name in employee_names:
            print(f"正在查找与姓名 '{employee_name}' 相关的PDF文件...")
            pdf_file = find_pdf_files(self.pdf_dir, employee_name)

            if pdf_file:
                # 将姓名和找到的PDF文件路径存入字典
                name_pdf_mapping[employee_name] = pdf_file
                match_files.append(pdf_file)
            else:
                print(f"未找到与姓名 '{employee_name}' 相关的PDF文件...")
                name_pdf_mapping[employee_name] = None

        # 判断是否找到
        total_files = len(match_files)  # 总文件数
        if total_files == 0:
            print("没有找到符合条件的PDF文件")
            return

        para1 = doc.add_paragraph()
        parent_number = ""
        prefix = ""

        run1 = para1.add_run(f"{parent_number}xxx分公司社保")
        run1.bold = True
        run1.font.size = Pt(16)
        run1.font.color.rgb = RGBColor(0, 0, 0)  # 设置字体颜色为黑色
        # 设置字体为宋体
        run1.font.name = '宋体'
        # 设置东亚字体（中文）为宋体
        run1._r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        para1.style = 'Heading 2'
        # 插入序号和姓名，并更新进度条
        i = 0
        for match_name, match_pdf_path in name_pdf_mapping.items():
            # 插入序号和姓名
            i += 1
            if prefix != "":
                seq_number = ".".join(map(str, prefix))
                prefix[-1] += 1
                seq_number = seq_number + " "
            else:
                seq_number = ""

            if match_pdf_path:
                exits_name = match_name
            else:
                exits_name =  f"{match_name}(无对应文件)"

            para2 = doc.add_paragraph()
            run2 = para2.add_run(f"{seq_number}{exits_name}")
            run2.bold = True
            run2.font.size = Pt(14)
            run2.font.color.rgb = RGBColor(0, 0, 0)  # 设置字体颜色为黑色
            # 设置字体为宋体
            run2.font.name = '宋体'
            # 设置东亚字体（中文）为宋体
            run2._r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            para2.style = 'Heading 3'

            # 插入 PDF 文件的所有页作为图片
            if match_pdf_path:
                self.insert_pdf_as_image(doc, match_pdf_path, extract_chinese(exits_name))
            else:
                para2.add_run("\n")  # 换行
                print(f"未找到 PDF 文件: {match_pdf_path}")

        print("所有PDF文件处理完成")

        # 保存最终的 Word 文件
        current_time = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_word_file = os.path.join(os.getcwd(), 'output', f'socail_documents_{current_time}.docx')
        doc.save(output_word_file)
        # 打开文件
        os.startfile(output_word_file)

    # 将 PDF 转换为图片，并插入到 Word
    def insert_pdf_as_image(self, doc, pdf_filepath, keyword):
        try:
            # 打开PDF文件
            doc_pdf = fitz.open(pdf_filepath)
            num_pages = doc_pdf.page_count
            print(f"Total pages in PDF: {num_pages}")

            # 设置DPI来提高图像分辨率
            dpi = 300  # 可根据需求调整
            target_width_inch = 5.91  # 目标宽度：15厘米，即5.91英寸
            target_width_px = target_width_inch * 96  # 转换为像素

            for page_num in range(num_pages):
                # 加载PDF页面
                page = doc_pdf.load_page(page_num)
                print(f"Converting page {page_num + 1} to image...")

                PdfProcessUtils.draw_rect_name(page, keyword)
                PdfProcessUtils.draw_rect_unit(page)

                # 以下为裁剪掉文档空白区域，先获取页面内容的边界（例如，文本块的最大/最小坐标）
                text_dict = page.get_text("dict")

                # 打印页面基本信息以便调试
                print(f"Processing page {page_num + 1} of {num_pages}")
                print(f"Page size: {page.rect}")

                if text_dict["blocks"]:  # 检查是否有文本块
                    min_y = float('inf')
                    max_y = float('-inf')

                    for block in text_dict["blocks"]:
                        if block['type'] == 0:  # 0表示文本块
                            for line in block['lines']:
                                for span in line['spans']:
                                    min_y = min(min_y, span['bbox'][1])
                                    max_y = max(max_y, span['bbox'][3])

                    # 打印识别到的文本区域
                    print(f"Text block min_y: {min_y}, max_y: {max_y}")

                    # 计算页面内容的高度
                    page_content_height = max_y - min_y
                    page_height = page.rect.height
                    blank_threshold = 0.3  # 假设页面底部 30% 是空白区域

                    # 打印空白区域的比例
                    print(f"Content height: {page_content_height}, Page height: {page_height}")
                    print(f"Blank area ratio: {(page_height - page_content_height) / page_height:.2f}")

                    # 如果页面的内容高度占比小于阈值，则裁剪掉底部的空白区域
                    if page_content_height / page_height < (1 - blank_threshold):
                        print("Content occupies less than 70% of the page, cropping bottom area.")
                        # pix = page.get_pixmap(clip=fitz.Rect(0, min_y-50, page.rect.width, max_y+50))
                        pix = page.get_pixmap(clip=fitz.Rect(0, 0, page.rect.width, max_y + 50),dpi=dpi)
                    else:
                        print("Content occupies more than 70% of the page, no cropping.")
                        pix = page.get_pixmap(dpi=dpi)  # 否则，渲染整个页面
                else:
                    # 如果页面没有文本块，直接渲染整个页面
                    print("No text blocks found, rendering the entire page.")
                    pix = page.get_pixmap(dpi=dpi)
                # 获取PDF页面的Pixmap对象
                # pix = page.get_pixmap(dpi=dpi)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                # 获取图像的宽高
                img_width, img_height = img.size

                # 锁定纵横比，计算相应的高度
                scale = target_width_px / img_width
                target_height_px = img_height * scale

                # 输出缩放后的图像尺寸
                print(f"Page {page_num + 1} image size: {target_width_px:.0f}x{target_height_px:.0f}")

                # 将图像保存到内存
                img_byte_arr = BytesIO()
                img.save(img_byte_arr, format='PNG')
                img_byte_arr.seek(0)  # 重置游标到开始位置

                # 将图片插入到Word文档
                doc.add_picture(img_byte_arr, width=Inches(target_width_inch), height=Inches(target_height_px / 96))

                # 清理内存
                gc.collect()

            print("PDF to Word conversion complete.")

        except Exception as e:
            print(f"Error: {e}")
            gc.collect()  # 出现错误时进行垃圾回收


